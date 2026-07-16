from pathlib import Path
import sys
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")
source = Path(r"D:\Dowload\CẦN-SỬA-TỔNG-QUAN.docx")
document = Document(source)

for index, paragraph in enumerate(document.paragraphs, start=1):
    text = paragraph.text.strip()
    if text:
        print(f"P{index} [{paragraph.style.name}] {text}")

for index, table in enumerate(document.tables, start=1):
    print(f"\nTABLE {index}")
    for row in table.rows:
        print(" | ".join(cell.text.strip().replace("\n", " / ") for cell in row.cells))
