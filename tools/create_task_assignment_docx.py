from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8")

OUTPUT = Path(r"E:\Code\Dự án WMS\REPO git\WMS-\docs\PHAN-CONG-CONG-VIEC-KHAI-GIANG.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = "1F2937"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def fix_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    for index, text in enumerate(headers):
        cell = header.cells[index]
        shade(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, 9, True, DARK_BLUE)
    for row_values in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, 9, False, INK)
            if index in (0, len(row_values) - 1):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fix_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run, 11, False, INK)
    return paragraph


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_callout(document, label, text):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    fix_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, 10, True, DARK_BLUE)
    text_run = p.add_run(text)
    set_run_font(text_run, 10, False, INK)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


document = Document()
section = document.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
configure_styles(document)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_run = header.add_run("WMS | PHÂN CÔNG CÔNG VIỆC")
set_run_font(header_run, 8, True, "6B7280")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("Nội bộ nhóm dự án WMS")
set_run_font(footer_run, 8, False, "6B7280")

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
title_run = title.add_run("PHÂN CÔNG XỬ LÝ CÁC HẠNG MỤC WMS")
set_run_font(title_run, 20, True, DARK_BLUE)

subtitle = document.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
subtitle_run = subtitle.add_run("Căn cứ tài liệu “CẦN SỬA TỔNG QUAN” | Phân công cho Trần Quang Khải và Lê Hoàng Giang")
set_run_font(subtitle_run, 10, False, "6B7280")

document.add_heading("1. Mục tiêu và nguyên tắc phân công", level=1)
paragraph = document.add_paragraph()
paragraph.paragraph_format.space_after = Pt(6)
run = paragraph.add_run(
    "Mục tiêu là chuyển WMS từ mô hình xem dữ liệu toàn hệ thống sang vận hành theo hub, "
    "đồng thời chỉnh các luồng pallet, đơn nhập/xuất, xe, chuyến và nhật ký theo yêu cầu tổng quan."
)
set_run_font(run, 11, False, INK)
add_callout(
    document,
    "Phụ thuộc cần chốt",
    "Duy Anh cần quyết định hướng kiến trúc trước khi triển khai sâu: (1) làm lại mô hình dữ liệu và service theo hub, "
    "hoặc (2) giữ mô hình admin tổng và mở rộng dần role Admin hub, Quản lý kho, Scanner, Driver. "
    "Các đầu việc dưới đây ưu tiên cách làm có thể áp dụng cho cả hai hướng.",
)
add_bullet(document, "Trần Quang Khải sở hữu phần dữ liệu, backend service, rule phân quyền và migration cần thiết.")
add_bullet(document, "Lê Hoàng Giang sở hữu màn hình, trải nghiệm vận hành, biểu mẫu tạo mới và kiểm thử giao diện.")
add_bullet(document, "Mọi thay đổi API hoặc schema phải được Khải công bố trước để Giang tích hợp giao diện theo cùng contract.")

document.add_heading("2. Phân công Trần Quang Khải", level=1)
khai_rows = [
    ("KHAI-01", "Phân quyền theo hub", "Thiết kế phạm vi dữ liệu theo hub/location: nhân viên chỉ xem hub được gán; Manager/Admin có phạm vi theo role.", "Schema/claim phạm vi hub, policy backend, API không lộ dữ liệu hub khác.", "P0"),
    ("KHAI-02", "Service và DB", "Rà soát entity, migration và service bị ảnh hưởng bởi hub scope; cập nhật truy vấn Locations, Zones, Pallets, Trips, Orders.", "Migration có rollback; service lọc hub thống nhất; tài liệu API cập nhật.", "P0"),
    ("KHAI-03", "Định danh tự sinh", "Chuẩn hóa quy tắc sinh ID cho pallet, xe và nhân viên; chốt format cùng Duy Anh.", "API tạo mới không nhận ID do người dùng tự nhập; kiểm tra không trùng mã.", "P1"),
    ("KHAI-04", "Nghiệp vụ đơn và chuyến", "Đổi nghĩa Inbound supplier thành hub xuất phát; loại customer ở Outbound; lọc lịch chuyến theo location của người dùng.", "DTO/API mới, dữ liệu cũ được chuyển đổi hoặc có tương thích ngược.", "P1"),
    ("KHAI-05", "Route và audit", "Đưa routing rules về service chạy ngầm; tạo API tra tuyến theo điểm đi/đến. Phân loại audit theo đối tượng xe, pallet, đơn, bao.", "API route finder; audit có filter type/table và phân trang.", "P2"),
]
add_table(document, ["Mã", "Hạng mục", "Việc thực hiện", "Bàn giao", "Ưu tiên"], khai_rows, [720, 1260, 3000, 3300, 1080])

document.add_heading("3. Phân công Lê Hoàng Giang", level=1)
giang_rows = [
    ("GIANG-01", "Giao diện theo hub", "Hiển thị hub đang làm việc; ẩn dữ liệu, menu và hành động không thuộc quyền/hub của người dùng.", "UI không hiển thị dữ liệu hub khác; trạng thái không quyền rõ ràng.", "P0"),
    ("GIANG-02", "Zone và pallet", "Lọc zone theo hub của nhân viên; thêm tìm pallet theo ID; làm form tạo pallet không nhập ID thủ công.", "Màn tìm kiếm và tạo pallet; mã trả về từ server hiển thị sau khi tạo.", "P1"),
    ("GIANG-03", "Inbound và Outbound", "Đổi nhãn/biểu mẫu Inbound thành hub xuất phát; bỏ customer khỏi UI Outbound; đồng bộ validation theo API mới.", "Form và danh sách phản ánh đúng thuật ngữ nghiệp vụ.", "P1"),
    ("GIANG-04", "Fleet và Trip", "Làm màn tạo xe với ID tự sinh; lịch chuyến chỉ hiện dữ liệu location đang thao tác.", "Form tạo xe, bảng lịch chuyến có lọc location và trạng thái tải/rỗng/lỗi.", "P1"),
    ("GIANG-05", "Routing và Audit", "Thay danh sách routing rules rối bằng form chọn điểm đi/đến rồi hiện tuyến; bổ sung bộ lọc audit theo loại log.", "Màn tìm tuyến và màn audit có bộ lọc dễ dùng.", "P2"),
    ("GIANG-06", "Kiểm thử UI", "Cập nhật kịch bản kiểm thử cho role/hub, mã tự sinh và toàn bộ luồng mới; ghi nhận lỗi tái hiện được.", "Checklist test, ảnh/chứng cứ và danh sách lỗi bàn giao.", "P0"),
]
add_table(document, ["Mã", "Hạng mục", "Việc thực hiện", "Bàn giao", "Ưu tiên"], giang_rows, [720, 1260, 3000, 3300, 1080])

document.add_page_break()
document.add_heading("4. Các bước thực hiện chi tiết", level=1)
add_callout(
    document,
    "Cách dùng",
    "Mỗi đầu việc chỉ được chuyển sang trạng thái hoàn thành khi đủ bước thực hiện, có commit/PR tương ứng và có bằng chứng kiểm thử. "
    "Nếu phát hiện thay đổi phụ thuộc đầu việc của người còn lại, phải ghi rõ API contract hoặc migration cần chờ.",
)

document.add_heading("4.1 Trần Quang Khải - Backend, database và phân quyền", level=2)
khai_details = [
    ("KHAI-01 | Phân quyền theo hub", [
        "Liệt kê toàn bộ role hiện có và role mục tiêu: Admin tổng, Admin hub, Quản lý kho, Scanner và Driver.",
        "Chốt quy tắc phạm vi dữ liệu: mỗi tài khoản được gán hub/location nào, role nào được xem nhiều hub và role nào chỉ xem hub hiện tại.",
        "Thiết kế migration bổ sung quan hệ quyền với hub/location; thêm index và ràng buộc để tránh gán trùng hoặc dữ liệu mồ côi.",
        "Tạo claim/policy và hàm lọc dùng chung trong service; áp dụng cho API đọc, tạo, sửa, xóa thay vì chỉ ẩn UI.",
        "Tạo dữ liệu kiểm thử cho ít nhất hai hub; dùng API trực tiếp để chứng minh tài khoản hub A không đọc/sửa được dữ liệu hub B.",
    ]),
    ("KHAI-02 | Rà soát service và DB", [
        "Lập danh sách endpoint, entity và service bị ảnh hưởng: Locations, Zones, Pallets, Employees, Trips, Inbound, Outbound, Reservations và Audit.",
        "Xác định khóa location/hub nguồn của từng bản ghi và cách xử lý dữ liệu chưa có hub rõ ràng.",
        "Viết migration theo từng thay đổi nhỏ, có kiểm tra dữ liệu trước/sau và phương án rollback.",
        "Cập nhật query service để luôn nhận context hub từ tài khoản hiện tại, không nhận hub tùy ý từ client nếu role không cho phép.",
        "Công bố API contract, mã lỗi và dữ liệu seed cho Giang trước khi giao diện tích hợp.",
    ]),
    ("KHAI-03 | ID tự sinh", [
        "Chốt format ID với Duy Anh: prefix, thời gian/sequence, độ dài tối đa và ví dụ cho Pallet, Car, Employee.",
        "Bổ sung generator ở server; client không được gửi ID khi tạo mới.",
        "Thiết lập unique constraint/index tương ứng và xử lý retry khi có va chạm mã.",
        "Cập nhật response API để trả mã vừa sinh; ghi audit cho thao tác tạo mới.",
        "Viết test tạo đồng thời hoặc lặp nhiều lần để kiểm tra không trùng ID.",
    ]),
    ("KHAI-04 | Đơn và chuyến", [
        "Đối chiếu schema hiện tại của Inbound, Outbound và Trip với nghiệp vụ hub xuất phát/hub đích.",
        "Thay trường supplier của Inbound bằng location/hub xuất phát; loại customer của Outbound nếu không còn dùng trong nghiệp vụ.",
        "Cập nhật DTO, validation, migration và dữ liệu seed; giữ mapping tương thích cho dữ liệu cũ nếu cần.",
        "Áp dụng lọc trip schedule theo location của người dùng và role điều phối.",
        "Kiểm thử tạo, sửa, xuất và nhận đơn ở hai hub khác nhau.",
    ]),
    ("KHAI-05 | Routing và audit", [
        "Chuyển routing rules thành service nội bộ; xác định đầu vào là hub đi và hub đến.",
        "Tạo API tra tuyến trả về next hop/chuỗi tuyến tối giản, không trả danh sách rule thô cho người dùng thường.",
        "Chuẩn hóa nhóm audit: xe, pallet, bao hàng, đơn nhập, đơn xuất, tài khoản và phân quyền.",
        "Bổ sung filter theo nhóm log, thời gian, hub và phân trang để tránh tải toàn bộ audit log.",
        "Kiểm thử route finder và audit filter với dữ liệu lớn hơn seed mẫu.",
    ]),
]
for heading, steps in khai_details:
    document.add_heading(heading, level=3)
    for step in steps:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(step)
        set_run_font(run, 10, False, INK)

document.add_heading("4.2 Lê Hoàng Giang - Giao diện và kiểm thử vận hành", level=2)
giang_details = [
    ("GIANG-01 | Giao diện theo hub", [
        "Nhận API contract phạm vi hub từ Khải; bổ sung trạng thái hub hiện tại vào layout/topbar.",
        "Lọc menu, danh sách, bộ chọn và hành động theo role trả về từ server; không chỉ dựa vào dữ liệu local của frontend.",
        "Tạo trạng thái rỗng/không quyền/lỗi tải dữ liệu rõ ràng cho từng màn hình.",
        "Kiểm thử đăng nhập bằng ít nhất Manager, WarehouseStaff và Driver để xác nhận menu đúng role.",
    ]),
    ("GIANG-02 | Zone và pallet", [
        "Cập nhật màn Zones để chỉ tải zone của hub hiện tại theo API mới.",
        "Thêm ô tìm kiếm pallet theo ID, hiển thị hub, zone, trạng thái và thông báo không tìm thấy.",
        "Tạo form Add Pallet không có ô nhập ID; người dùng chỉ chọn zone/capacity, sau đó nhận ID do server trả về.",
        "Kiểm thử tạo pallet, tìm theo ID và thử truy cập pallet ở hub khác.",
    ]),
    ("GIANG-03 | Inbound và Outbound", [
        "Đổi toàn bộ label, placeholder, table header và validation của Inbound từ supplier sang hub xuất phát.",
        "Loại customer khỏi form/danh sách Outbound theo DTO mới; kiểm tra các trang dashboard hoặc detail có dùng trường này.",
        "Cập nhật API client, type TypeScript và thông báo lỗi theo response backend.",
        "Thực hiện kiểm thử nhập, chia chọn, giữ hàng và xuất hàng bằng dữ liệu demo.",
    ]),
    ("GIANG-04 | Fleet và Trip", [
        "Thêm form tạo xe: loại xe, sức chứa; không cho nhập car ID.",
        "Sau khi tạo thành công, hiển thị ID xe do server trả về và tải lại danh sách.",
        "Cập nhật Trip Scheduling để mặc định lọc theo location/hub; chỉ role được cấp quyền mới đổi phạm vi xem.",
        "Kiểm thử trạng thái loading, không có chuyến, lỗi quyền và lịch của hai hub.",
    ]),
    ("GIANG-05 | Routing và Audit", [
        "Thay bảng routing rules trực tiếp bằng form chọn điểm đi và điểm đến.",
        "Gọi API route finder, hiển thị kết quả tuyến theo dạng dễ quét: điểm đi, các next hop, điểm đến.",
        "Thêm filter audit theo nhóm đối tượng, hub, khoảng thời gian và trạng thái tải trang.",
        "Kiểm thử URL, reload trang, dữ liệu trống và lỗi API cho cả hai màn hình.",
    ]),
    ("GIANG-06 | Kiểm thử UI", [
        "Cập nhật file kịch bản kiểm thử theo role/hub, ID tự sinh, luồng inbound/outbound, xe, pallet, trip, routing và audit.",
        "Chạy test smoke sau mỗi API contract mới; ghi mã test, bước tái hiện, ảnh màn hình và kết quả thực tế.",
        "Tạo danh sách lỗi theo mức độ P0/P1/P2, chỉ rõ endpoint hoặc màn hình liên quan và điều kiện tái hiện.",
        "Xác nhận lại các tiêu chí nghiệm thu chung trước khi đóng sprint.",
    ]),
]
for heading, steps in giang_details:
    document.add_heading(heading, level=3)
    for step in steps:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(step)
        set_run_font(run, 10, False, INK)

document.add_heading("5. Trình tự phối hợp", level=1)
steps = [
    "Duy Anh chốt hướng kiến trúc và format ID trước khi bắt đầu các hạng mục P0.",
    "Khải chốt migration, API contract và policy quyền; gửi mẫu response/error cho Giang.",
    "Giang tích hợp UI theo contract, không hard-code dữ liệu hub hoặc ID.",
    "Cả hai kiểm thử chéo bằng ít nhất Admin/Manager, nhân viên kho và Driver ở hai hub khác nhau.",
    "Chỉ merge khi API không lộ dữ liệu chéo hub, các ID mới tự sinh và kịch bản kiểm thử P0 đạt.",
]
for step in steps:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(step)
    set_run_font(run, 11, False, INK)

document.add_heading("6. Tiêu chí nghiệm thu chung", level=1)
for item in [
    "Tài khoản không có quyền hoặc không thuộc hub không thể xem/sửa dữ liệu ngoài phạm vi, kể cả gọi API trực tiếp.",
    "Pallet, xe và nhân viên mới có ID do server sinh; UI chỉ hiển thị ID sau khi tạo thành công.",
    "Inbound thể hiện hub xuất phát; Outbound không còn trường customer nếu không còn nhu cầu nghiệp vụ.",
    "Trip schedule, zone, pallet và routing hiển thị đúng theo location/hub đang thao tác.",
    "Audit có thể lọc theo nhóm đối tượng để người vận hành không phải xem toàn bộ log hỗn hợp.",
]:
    add_bullet(document, item)

document.core_properties.title = "Phân công xử lý các hạng mục WMS"
document.core_properties.subject = "Phân công công việc cho Trần Quang Khải và Lê Hoàng Giang"
document.core_properties.author = "Nhóm WMS"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
document.save(OUTPUT)
print(OUTPUT)
